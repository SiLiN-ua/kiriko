"""
Kiriko — Flask application
"""
import os
import io
import re
import json
import uuid
import time
import hashlib
import asyncio
import threading
import functools
from collections import deque
from datetime import datetime
from flask import (Flask, render_template, request, redirect,
                   url_for, session, Response, jsonify, send_file)
from flask_cors import CORS
from dotenv import load_dotenv

# Ensure the main thread always has an event loop (needed for tools.py
# which calls asyncio.get_event_loop().run_until_complete() for Telegram bots).
try:
    asyncio.get_event_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

from core.agent import run_agent_stream

app = Flask(__name__)
CORS(app)

# ── History storage ───────────────────────────────────────────────────────────

DATA_DIR    = os.path.join(os.path.dirname(__file__), 'data')
HISTORY_FILE = os.path.join(DATA_DIR, 'history.json')
_history_lock = threading.Lock()

def _ensure_data():
    os.makedirs(DATA_DIR, exist_ok=True)
    if not os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump([], f)

def _load_history():
    _ensure_data()
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

def _save_history(history):
    _ensure_data()
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def _upsert_conversation(session_id, msgs):
    """Save or update a conversation keyed by session_id."""
    with _history_lock:
        history = _load_history()

        title = 'Беседа'
        for m in msgs:
            if m.get('role') == 'user':
                content = m.get('content', '')
                if isinstance(content, str):
                    title = content[:80]
                elif isinstance(content, list):
                    for block in content:
                        if isinstance(block, dict) and block.get('type') == 'text':
                            title = block.get('text', '')[:80]
                            break
                break

        now = datetime.now().isoformat(timespec='seconds')
        user_turns = len([m for m in msgs if m.get('role') == 'user'
                          and isinstance(m.get('content'), str)])

        for i, h in enumerate(history):
            if h.get('session_id') == session_id:
                history[i].update(timestamp=now, title=title,
                                  msg_count=user_turns, messages=msgs)
                _save_history(history)
                return

        entry = {
            'id':         str(uuid.uuid4()),
            'session_id': session_id,
            'timestamp':  now,
            'title':      title,
            'msg_count':  user_turns,
            'messages':   msgs,
        }
        history.insert(0, entry)
        _save_history(history[:300])

PASSWORD = os.getenv('APP_PASSWORD', '')
# Stable secret key derived from the app password
app.secret_key = hashlib.sha256(f'kiriko-{PASSWORD}'.encode()).digest()


# ── Rate limiter (per session, sliding window) ────────────────────────────────
# Limits: 5 requests per minute, 20 per hour per session

_rl_lock   = threading.Lock()
_rl_store  = {}   # session_id → {'min': deque, 'hour': deque}

RATE_PER_MIN  = int(os.getenv('RATE_PER_MIN',  '5'))
RATE_PER_HOUR = int(os.getenv('RATE_PER_HOUR', '20'))

def _check_rate_limit(sid: str):
    """
    Returns (allowed: bool, reason: str).
    Sliding window: prune timestamps older than the window, then count.
    """
    now = time.monotonic()
    with _rl_lock:
        if sid not in _rl_store:
            _rl_store[sid] = {'min': deque(), 'hour': deque()}
        buckets = _rl_store[sid]

        # Prune expired
        while buckets['min']  and now - buckets['min'][0]  > 60:
            buckets['min'].popleft()
        while buckets['hour'] and now - buckets['hour'][0] > 3600:
            buckets['hour'].popleft()

        if len(buckets['min']) >= RATE_PER_MIN:
            wait = int(60 - (now - buckets['min'][0])) + 1
            return False, f'Забагато запитів — зачекайте {wait} с (ліміт {RATE_PER_MIN}/хв)'
        if len(buckets['hour']) >= RATE_PER_HOUR:
            wait = int(3600 - (now - buckets['hour'][0])) // 60 + 1
            return False, f'Годинний ліміт вичерпано — зачекайте ~{wait} хв ({RATE_PER_HOUR}/год)'

        buckets['min'].append(now)
        buckets['hour'].append(now)
        return True, ''


# ── Auth helper ───────────────────────────────────────────────────────────────

def require_auth(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('authenticated'):
            return redirect(url_for('splash'))
        return f(*args, **kwargs)
    return decorated


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return redirect(url_for('splash'))


@app.route('/splash')
def splash():
    authenticated = bool(session.get('authenticated'))
    return render_template('splash.html', authenticated=authenticated)


@app.route('/api/auth', methods=['POST'])
def auth():
    data = request.get_json(silent=True) or {}
    if data.get('password', '') == PASSWORD:
        session['authenticated'] = True
        session.permanent = True
        return jsonify({'ok': True})
    return jsonify({'ok': False, 'error': 'Неверный пароль'}), 401


@app.route('/chat')
@require_auth
def chat():
    return render_template('chat.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('splash'))


@app.route('/api/chat', methods=['POST'])
@require_auth
def api_chat():
    data       = request.get_json(silent=True) or {}
    messages   = data.get('messages', [])
    session_id = data.get('session_id', '')
    if not messages:
        return jsonify({'error': 'No messages'}), 400

    sid = session.get('_id') or request.remote_addr or 'default'
    allowed, reason = _check_rate_limit(sid)
    if not allowed:
        return jsonify({'error': reason}), 429

    def generate():
        try:
            for event_json in run_agent_stream(messages):
                yield f'data: {event_json}\n\n'
                if session_id:
                    try:
                        evt = json.loads(event_json)
                        if evt.get('type') == 'done' and evt.get('messages'):
                            _upsert_conversation(session_id, evt['messages'])
                    except Exception:
                        pass
        except Exception as exc:
            yield f'data: {json.dumps({"type": "error", "text": str(exc)})}\n\n'
            yield f'data: {json.dumps({"type": "done", "messages": messages})}\n\n'

    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive',
        },
    )


@app.route('/api/history', methods=['GET'])
@require_auth
def api_history():
    history = _load_history()
    return jsonify([{
        'id':        h['id'],
        'timestamp': h.get('timestamp', ''),
        'title':     h.get('title', 'Беседа'),
        'msg_count': h.get('msg_count', 0),
    } for h in history])


@app.route('/api/history/<conv_id>', methods=['GET'])
@require_auth
def api_history_get(conv_id):
    history = _load_history()
    conv = next((h for h in history if h['id'] == conv_id), None)
    if not conv:
        return jsonify({'error': 'not found'}), 404
    return jsonify(conv)


@app.route('/api/history/<conv_id>', methods=['DELETE'])
@require_auth
def api_history_delete(conv_id):
    history = _load_history()
    history = [h for h in history if h['id'] != conv_id]
    _save_history(history)
    return jsonify({'ok': True})


@app.route('/api/history', methods=['DELETE'])
@require_auth
def api_history_clear_all():
    """Delete ALL saved conversations."""
    _save_history([])
    return jsonify({'ok': True})


# ── File upload → content blocks ──────────────────────────────────────────────

@app.route('/api/upload', methods=['POST'])
@require_auth
def api_upload():
    """
    Accepts a file upload and returns Claude API content blocks.
    Images  → [{"type":"image", "source":{"type":"base64",...}}]
    PDF     → text blocks if text PDF, or image blocks if scanned
    DOCX    → [{"type":"text", "text":"extracted text"}]
    """
    if 'file' not in request.files:
        return jsonify({'error': 'no file'}), 400

    f    = request.files['file']
    name = f.filename or 'file'
    data = f.read()
    ext  = os.path.splitext(name)[1].lower()

    # ── DOCX / DOC ────────────────────────────────────────────────────────────
    if ext in ('.docx', '.doc'):
        if ext == '.doc':
            try:
                import subprocess, tempfile
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                    tmp.write(data)
                    tmp_path = tmp.name
                result = subprocess.run(['antiword', tmp_path], capture_output=True, timeout=30)
                os.unlink(tmp_path)
                text = result.stdout.decode('utf-8', errors='ignore').strip()
                if not text:
                    text = '[DOC: не вдалося витягти текст]'
                blocks = [{'type': 'text', 'text': f'[Файл: {name}]\n{text}'}]
                return jsonify({'blocks': blocks})
            except Exception as e:
                return jsonify({'error': f'Помилка читання .doc: {e}'}), 500
        try:
            from docx import Document as DocxDoc
            doc  = DocxDoc(io.BytesIO(data))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip() and cell.text.strip() not in parts:
                            parts.append(cell.text.strip())
            text = '\n'.join(parts) if parts else '[DOCX: порожній або не вдалося зчитати вміст]'
            blocks = [{'type': 'text', 'text': f'[Файл: {name}]\n{text}'}]
            return jsonify({'blocks': blocks})
        except Exception as e:
            return jsonify({'error': f'Помилка читання DOCX: {e}'}), 500

    # ── PDF ───────────────────────────────────────────────────────────────────
    if ext == '.pdf':
        try:
            import pymupdf
            doc  = pymupdf.open(stream=data, filetype='pdf')
            text = ''.join(page.get_text() for page in doc).strip()

            # Text PDF — send extracted text
            if len(text) > 150:
                # Limit to ~12000 chars to stay within token budget
                if len(text) > 12000:
                    text = text[:12000] + '\n[...текст скорочено...]'
                blocks = [{'type': 'text', 'text': f'[Файл: {name}]\n{text}'}]
                return jsonify({'blocks': blocks})

            # Scanned PDF — render pages as images (max 10 pages)
            import base64
            blocks = [{'type': 'text', 'text': f'[Файл: {name} — скан, сторінок: {len(doc)})'}]
            for i, page in enumerate(doc):
                if i >= 10:
                    blocks.append({'type': 'text', 'text': f'[...ще {len(doc)-10} сторінок не показано]'})
                    break
                pix     = page.get_pixmap(dpi=150)
                img_b64 = base64.b64encode(pix.tobytes('png')).decode()
                blocks.append({
                    'type': 'image',
                    'source': {
                        'type':       'base64',
                        'media_type': 'image/png',
                        'data':       img_b64,
                    }
                })
            return jsonify({'blocks': blocks})
        except Exception as e:
            return jsonify({'error': f'PDF error: {e}'}), 500

    # ── Image (already handled client-side, but accept server-side too) ────────
    mime_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png',  '.gif':  'image/gif',
        '.webp': 'image/webp',
    }
    if ext in mime_map:
        import base64
        b64 = base64.b64encode(data).decode()
        blocks = [{
            'type': 'image',
            'source': {'type': 'base64', 'media_type': mime_map[ext], 'data': b64}
        }]
        return jsonify({'blocks': blocks})

    return jsonify({'error': f'Unsupported file type: {ext}'}), 400


# ── DOCX export ───────────────────────────────────────────────────────────────

@app.route('/api/export/docx', methods=['POST'])
@require_auth
def api_export_docx():
    """
    Accepts JSON { "text": "<plain report text>" }
    Formatting rules:
      - Times New Roman 14pt, justified
      - BOLD: first line (ПІБ), section headers, relative headers
      - NORMAL: all other content
      - URLs → clickable hyperlinks (blue underlined)
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants

    data = request.get_json(silent=True) or {}
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'error': 'no text'}), 400

    # ── Логіка жирного: "Мітка: значення" ────────────────────────────────
    # Рядки що починаються з дати (записи роботи/освіти) — повністю НЕ жирний
    _DATE_LINE_RE = re.compile(r'^\s*(\d{2}[./]\d{2}[./]\d{4}|\d{4})\s*[-––—]')

    def _get_line_segments(raw_line: str, is_first: bool):
        """
        Повертає список (text, is_bold).
        - Перший рядок: ФИО (до першої коми) = жирний, решта = ні
        - Рядки що починаються з дати = повністю НЕ жирний
        - "Мітка: значення" = мітка жирна, значення НЕ жирне
        - Рядки без двокрапки = НЕ жирний
        """
        line = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', raw_line)
        line = re.sub(r'^#+\s*', '', line)
        if not line.strip():
            return []

        # Перший рядок: ФИО до першої коми — жирний
        if is_first:
            comma_idx = line.find(',')
            if comma_idx > 0:
                return [(line[:comma_idx], True), (line[comma_idx:], False)]
            return [(line, True)]

        # Рядки що починаються з дати (записи освіти/роботи) — НЕ жирний
        if _DATE_LINE_RE.match(line):
            return [(line, False)]

        # Рядки з двокрапкою: мітка жирна, значення НЕ жирне
        colon_idx = line.find(':')
        if colon_idx > 0:
            label = line[:colon_idx + 1]   # включно з двокрапкою
            value = line[colon_idx + 1:]   # все після двокрапки
            if value.strip():
                return [(label, True), (value, False)]
            return [(line, True)]           # тільки мітка — повністю жирний

        # Звичайний рядок без двокрапки — НЕ жирний
        return [(line, False)]

    _URL_RE = re.compile(r'(https?://\S+)')

    def _add_hyperlink(paragraph, url: str):
        """Додає URL як клікабельне гіперпосилання (синій, підкреслений, НЕ жирний)."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)

        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')
        rPr.append(sz)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        # Гіперпосилання ніколи НЕ жирне
        b_off = OxmlElement('w:b')
        b_off.set(qn('w:val'), '0')
        rPr.append(b_off)

        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = url
        r.append(t)
        hl.append(r)
        paragraph._p.append(hl)

    def _fill_paragraph(p, raw_line: str, is_first: bool):
        """
        Заповнює абзац з розбивкою 'мітка: значення' (мітка жирна).
        URL всередині значення → клікабельне гіперпосилання.
        """
        segments = _get_line_segments(raw_line, is_first)
        for text_seg, is_bold in segments:
            url_parts = _URL_RE.split(text_seg)
            for part in url_parts:
                if not part:
                    continue
                if _URL_RE.fullmatch(part):
                    _add_hyperlink(p, part)   # URL — завжди НЕ жирний, клікабельний
                else:
                    run = p.add_run(part)
                    run.font.name      = 'Times New Roman'
                    run.font.size      = Pt(14)
                    run.font.bold      = is_bold
                    run.font.color.rgb = RGBColor(0, 0, 0)

    # ── Документ ──────────────────────────────────────────────────────────
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name      = 'Times New Roman'
    style.font.size      = Pt(14)
    style.font.bold      = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(1.4)
    pf.space_after  = Pt(1.4)
    pf.line_spacing = Pt(16.99)

    lines      = text.split('\n')
    first_done = False

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if not line.strip():
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            continue

        _fill_paragraph(p, line, not first_done)
        first_done = True

    # ── Ім'я файлу з першого рядка ────────────────────────────────────────
    first_line = next((l.strip() for l in lines if l.strip()), 'report')
    name_part  = first_line.split(',')[0].strip()
    safe_name  = re.sub(r'[^\w\s-]', '', name_part, flags=re.UNICODE).strip()
    safe_name  = re.sub(r'\s+', '_', safe_name) or 'osint_report'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f'{safe_name}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    app.run(
        host='127.0.0.1',
        port=5001,
        debug=False,
        threaded=True,    # multiple Flask threads — safe: asyncio in dedicated _tg_loop thread
        use_reloader=False,
    )
